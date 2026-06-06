"""
tests/test_resume_parser.py
Unit tests for resume text extraction.
"""

import sys
import os
import unittest
import io

sys.path.insert(0, os.path.dirname(os.path.dirname(os.path.abspath(__file__))))

from utils.resume_parser import extract_text_from_docx, extract_text


class TestResumeParser(unittest.TestCase):

    def test_unsupported_format(self):
        result = extract_text(b"dummy", "resume.txt")
        self.assertIn("Unsupported", result)

    def test_empty_bytes_pdf(self):
        result = extract_text(b"", "resume.pdf")
        # Should not crash; may return empty or error string
        self.assertIsInstance(result, str)

    def test_docx_extraction(self):
        """Create a minimal DOCX in memory and verify extraction."""
        from docx import Document
        doc = Document()
        doc.add_paragraph("Rupam Mukherjee")
        doc.add_paragraph("Skills: Python, Machine Learning, SQL")
        buf = io.BytesIO()
        doc.save(buf)
        buf.seek(0)
        text = extract_text_from_docx(buf.read())
        self.assertIn("Rupam", text)
        self.assertIn("Python", text)

    def test_filename_routing_pdf(self):
        """extract_text routes .pdf to PDF extractor (returns str even on bad bytes)."""
        result = extract_text(b"%PDF-bad", "cv.pdf")
        self.assertIsInstance(result, str)

    def test_filename_routing_docx(self):
        """extract_text routes .docx to DOCX extractor (returns error string on bad bytes)."""
        result = extract_text(b"bad_bytes", "cv.docx")
        self.assertIsInstance(result, str)


class TestSkillKeywords(unittest.TestCase):
    """Verify common skill keywords are preserved during extraction."""

    def test_skill_keywords_preserved_in_docx(self):
        from docx import Document
        doc = Document()
        doc.add_paragraph("Experience: Python, TensorFlow, Docker, React, SQL, AWS")
        buf = io.BytesIO()
        doc.save(buf)
        buf.seek(0)
        text = extract_text_from_docx(buf.read())
        for skill in ["Python", "TensorFlow", "Docker", "React", "SQL", "AWS"]:
            self.assertIn(skill, text)


if __name__ == "__main__":
    unittest.main()
