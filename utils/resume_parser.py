"""
utils/resume_parser.py
Extract text from PDF and DOCX resume files.
"""

import io
import PyPDF2
import pdfplumber
from docx import Document


def extract_text_from_pdf(file_bytes: bytes) -> str:
    """Extract text from a PDF file using pdfplumber (better formatting)."""
    text = ""
    try:
        with pdfplumber.open(io.BytesIO(file_bytes)) as pdf:
            for page in pdf.pages:
                page_text = page.extract_text()
                if page_text:
                    text += page_text + "\n"
    except Exception:
        # Fallback to PyPDF2
        try:
            reader = PyPDF2.PdfReader(io.BytesIO(file_bytes))
            for page in reader.pages:
                text += page.extract_text() + "\n"
        except Exception as e:
            text = f"[Error extracting PDF text: {e}]"
    return text.strip()


def extract_text_from_docx(file_bytes: bytes) -> str:
    """Extract text from a DOCX file."""
    text = ""
    try:
        doc = Document(io.BytesIO(file_bytes))
        for para in doc.paragraphs:
            text += para.text + "\n"
        # Also extract from tables
        for table in doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    text += cell.text + " "
                text += "\n"
    except Exception as e:
        text = f"[Error extracting DOCX text: {e}]"
    return text.strip()


def extract_text(file_bytes: bytes, filename: str) -> str:
    """Auto-detect file type and extract text."""
    filename_lower = filename.lower()
    if filename_lower.endswith(".pdf"):
        return extract_text_from_pdf(file_bytes)
    elif filename_lower.endswith(".docx"):
        return extract_text_from_docx(file_bytes)
    else:
        return "[Unsupported file format. Please upload PDF or DOCX.]"
